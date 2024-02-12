import os
import shutil
import pandas as pd
from docx import Document
import docx
from xml.etree.ElementTree import Element, SubElement, tostring
from xml.dom import minidom
from PIL import Image
import io
import magic
import random
import xml.etree.ElementTree as ET
import random
import string
import aspose.words as aw
from docx2txt import process
import docx2txt
from lien import title_doc
from lien import balise_docid
from lien import erreur_xml
from lien import cause_xml
from lien import solution_xml
from lien import details_xml
from lien import remarque_xml
from lien import contenue_xml
from lien import note_xml
from lien import procedure_xml
from lien import texte_xml
from lien import Détaillé_xml
from lien import consignes_xml
from test import replace_images_with_text

# Etape 1 : Charger le fichier xlsx et filtrer les documents 'Brouillon' et 'Publié'
xlsx_path = "NOUVELLE Base de connaissance\ImportArticleConnaissance_v12.xlsm"
df = pd.read_excel(xlsx_path)
print('Bien')
filtered_df = df[df['Statuts'].isin(['Brouillon', 'Publié'])]

# Etape 2 : Ajouter les noms des fichiers dans une liste 'Yes'
file_names_yes = filtered_df.iloc[:, 0].tolist()

# Etape 3 : Créer le répertoire 'FICHIER STATUT' et exporter les fichiers

replace_images_with_text('NOUVELLE Base de connaissance')

base_dir = "NOUVELLE Base de connaissance"
output_dir = os.path.join(base_dir, "FICHIER STATUT")

if not os.path.exists(output_dir):
    os.makedirs(output_dir)

for file_name in file_names_yes:
    source_path = os.path.join(base_dir, file_name)

    dest_path = os.path.join(output_dir, file_name)
    if os.path.exists(source_path):
        shutil.copy2(source_path, dest_path)
        print('C BON')

# Etape 4 : Exporter les fichiers de 'FICHIER STATUT' à 'BASE DE CONNAISSANCE' dans 'Code Python 2023'
output_code_dir = "base de connaissance"
print('YES')
output_base_dir = os.path.join(output_code_dir, "daitement")
print('O')
if not os.path.exists(output_base_dir):
    os.makedirs(output_base_dir)

for file_name in os.listdir(output_dir):
    source_path = os.path.join(output_dir, file_name)
    dest_path = os.path.join(output_base_dir, file_name)
    shutil.move(source_path, dest_path)




def convert_docx_to_text(docx_path):
    doc = Document(docx_path)
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return full_text




def convert_text_to_xml(text, xml_path):
    # Créer un élément racine
    root = Element("document")

    text = text.replace('Procédure Détaillé', 'Détaillé')

    text = text + 'NEWPORTTEXT'

    yes_st = text.split()

    title_doc(root, xml_path)

    balise_docid(root)

    erreur_xml(root, text)

    cause_xml(root, text)

    solution_xml(root, text)

    details_xml(root, text)

    note_xml(root, text)

    procedure_xml(root, text)

    Détaillé_xml(root, text)

    consignes_xml(root, text)

    remarque_xml(root, text)

    # Formatter l'arbre XML
    xml_content = minidom.parseString(tostring(root)).toprettyxml(indent="  ")

    # Enregistrer le résultat dans le fichier XML
    with open(xml_path, "w", encoding="utf-8") as xml_file:
        xml_file.write(xml_content)
    keywords = ["Erreur", "Cause", "Solution", "Détails", "ERREUR", "SOLUTION", "DÉTAILS", "Détail", "DÉTAIL", "CAUSE", "Solutions", "SOLUTIONS"]
    keywords_two = ['Procédure', 'PROCÉDURE', 'PROCEDURE', 'Procedure', "Note d'attention", "Note d'Attention", "Vue d'ensemble", "Vue d'Ensemble", "Procédure Détaillée"]
    for file_name in os.listdir(treated_dir):
        xml_path = os.path.join(treated_dir, file_name)
        
        # Ignorer les erreurs de décodage avec 'ignore'
        try:
            with open(xml_path, 'r', encoding='UTF-8', errors='ignore') as xml_file:
                content = xml_file.read()
        except Exception as e:
            print(f"Erreur lors de la lecture du fichier {file_name}: {e}")
            continue

        # Splitting the content into words
        words = content.split()

        for word in yes_st:
            if word in keywords_two:
                dest_path = os.path.join(mod_dir, file_name)
                break
            elif word in keywords:
                dest_path = os.path.join(kcs_dir, file_name)
                break
        else:
            dest_path = os.path.join(ref_dir, file_name)

        shutil.move(xml_path, dest_path)






docx_dir = os.path.join(output_base_dir, "Fichier DOCX")
treated_dir = os.path.join(output_base_dir, "Fichier TRIER")
if not os.path.exists(treated_dir):
    os.makedirs(treated_dir)
print('LLLLLL')
print('NO')

if not os.path.exists(docx_dir):
    os.makedirs(docx_dir)
if not os.path.exists(treated_dir):
    os.makedirs(treated_dir)




kcs_dir = os.path.join(output_base_dir, "Probleme Solution KCS")
ref_dir = os.path.join(output_base_dir, "Référence")
mod_dir = os.path.join(output_base_dir, "Modèle KCS")
print('NNNNN')

if not os.path.exists(kcs_dir):
    os.makedirs(kcs_dir)
if not os.path.exists(ref_dir):
    os.makedirs(ref_dir)
if not os.path.exists(mod_dir):
    os.makedirs(mod_dir)
print('YYYYYYYY')



















# Etape 6 : Convertir les fichiers DOCX en fichier XML
print('YES')

for file_name in os.listdir(output_base_dir):
    file_path = os.path.join(output_base_dir, file_name)
    dest_path = os.path.join(docx_dir, file_name)

    # Vérifier si le fichier est un fichier .docx
    if file_name.endswith(".docx") and os.path.isfile(file_path):
        shutil.move(file_path, dest_path)
for file_name in os.listdir(docx_dir):
    if file_name.endswith(".docx"):
        docx_path = os.path.join(docx_dir, file_name)
        xml_path = os.path.join(treated_dir, os.path.splitext(file_name)[0] + ".xml")

        try:
            doc = Document(docx_path)
            # Convertir le texte en XML
            content = convert_docx_to_text(docx_path)
            convert_text_to_xml(content, xml_path)
            print('Conversion réussie avec succès!')

        except Exception as e:
            print(f"Erreur lors de la conversion de {file_name}: {e}")

print('CCCCC')            

# Etape 8 : Analyser le contenu des fichiers XML et déplacer vers 'Probleme Solution KCS' ou 'Référence'

